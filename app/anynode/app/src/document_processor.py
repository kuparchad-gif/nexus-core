#!/usr/bin/env python3
"""
Document Processor for Viren Platinum Edition
Handles document parsing, analysis, and visualization
"""

import os
import json
import logging
import mimetypes
from typing import Dict, Any, List, Optional

# Configure logging
logger = logging.getLogger("DocumentProcessor")

class DocumentProcessor:
    """
    Processes various document types including PDF, DOCX, XLSX, and Visio
    """
    
    def __init__(self, upload_dir: str = "uploads", cache_dir: str = "cache"):
        """Initialize the document processor"""
        self.upload_dir = upload_dir
        self.cache_dir = cache_dir
        
        # Create directories if they don't exist
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(cache_dir, exist_ok=True)
        
        # Initialize document handlers
        self._init_handlers()
    
    def _init_handlers(self):
        """Initialize document type handlers"""
        # Register MIME types
        mimetypes.init()
        if not mimetypes.inited:
            mimetypes.init()
        
        # Add additional MIME types
        mimetypes.add_type('application/vnd.ms-visio.drawing', '.vsdx')
        mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
        mimetypes.add_type('application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', '.xlsx')
        
        # Define handlers
        self.handlers = {
            'application/pdf': self._process_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_xlsx,
            'application/vnd.ms-visio.drawing': self._process_visio,
            'text/plain': self._process_text,
            'text/csv': self._process_csv
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with processing results
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {"status": "error", "message": "File not found"}
        
        # Get file type
        mime_type, _ = mimetypes.guess_type(file_path)
        if not mime_type:
            logger.warning(f"Unknown file type for {file_path}")
            mime_type = 'application/octet-stream'
        
        # Get file handler
        handler = self.handlers.get(mime_type)
        if not handler:
            logger.warning(f"No handler for MIME type: {mime_type}")
            return {
                "status": "error",
                "message": f"Unsupported file type: {mime_type}",
                "file_path": file_path
            }
        
        # Process file
        try:
            result = handler(file_path)
            result["file_path"] = file_path
            result["mime_type"] = mime_type
            result["status"] = "success"
            return result
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                "status": "error",
                "message": f"Error processing document: {str(e)}",
                "file_path": file_path,
                "mime_type": mime_type
            }
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process a PDF document"""
        try:
            # Try to import PyPDF2
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                num_pages = len(reader.pages)
                
                # Extract text from first page
                first_page_text = reader.pages[0].extract_text()
                
                # Extract metadata
                metadata = reader.metadata
                
                return {
                    "type": "pdf",
                    "num_pages": num_pages,
                    "metadata": {
                        "title": metadata.title if metadata.title else "Unknown",
                        "author": metadata.author if metadata.author else "Unknown",
                        "subject": metadata.subject if metadata.subject else "Unknown",
                        "creator": metadata.creator if metadata.creator else "Unknown"
                    },
                    "preview": first_page_text[:500] + "..." if len(first_page_text) > 500 else first_page_text
                }
        except ImportError:
            logger.warning("PyPDF2 not installed, using fallback method")
            return {
                "type": "pdf",
                "message": "PDF processing requires PyPDF2 library",
                "preview": "PDF processing not available"
            }
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process a DOCX document"""
        try:
            # Try to import python-docx
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract text
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Count paragraphs, tables, etc.
            stats = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            }
            
            return {
                "type": "docx",
                "stats": stats,
                "preview": full_text[:500] + "..." if len(full_text) > 500 else full_text
            }
        except ImportError:
            logger.warning("python-docx not installed, using fallback method")
            return {
                "type": "docx",
                "message": "DOCX processing requires python-docx library",
                "preview": "DOCX processing not available"
            }
    
    def _process_xlsx(self, file_path: str) -> Dict[str, Any]:
        """Process an XLSX document"""
        try:
            # Try to import openpyxl
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            
            # Get sheet names
            sheet_names = workbook.sheetnames
            
            # Get preview of first sheet
            if sheet_names:
                sheet = workbook[sheet_names[0]]
                preview_data = []
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i >= 5:  # Only get first 5 rows
                        break
                    preview_data.append(list(row))
            
            return {
                "type": "xlsx",
                "sheets": sheet_names,
                "preview": preview_data if 'preview_data' in locals() else []
            }
        except ImportError:
            logger.warning("openpyxl not installed, using fallback method")
            return {
                "type": "xlsx",
                "message": "XLSX processing requires openpyxl library",
                "preview": "XLSX processing not available"
            }
    
    def _process_visio(self, file_path: str) -> Dict[str, Any]:
        """Process a Visio document"""
        # Visio processing requires specialized libraries
        return {
            "type": "visio",
            "message": "Visio processing requires specialized libraries",
            "preview": "Visio preview not available"
        }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process a text document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                # Count lines, words, characters
                lines = content.count('\n') + 1
                words = len(content.split())
                chars = len(content)
                
                return {
                    "type": "text",
                    "stats": {
                        "lines": lines,
                        "words": words,
                        "characters": chars
                    },
                    "preview": content[:500] + "..." if len(content) > 500 else content
                }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    return {
                        "type": "text",
                        "encoding": "latin-1",
                        "preview": content[:500] + "..." if len(content) > 500 else content
                    }
            except Exception as e:
                return {
                    "type": "text",
                    "message": f"Error reading text file: {str(e)}",
                    "preview": "Text preview not available"
                }
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process a CSV document"""
        try:
            import csv
            
            with open(file_path, 'r', newline='', encoding='utf-8') as file:
                reader = csv.reader(file)
                rows = []
                for i, row in enumerate(reader):
                    if i >= 5:  # Only get first 5 rows
                        break
                    rows.append(row)
                
                # Get header if available
                header = rows[0] if rows else []
                
                return {
                    "type": "csv",
                    "header": header,
                    "preview": rows
                }
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', newline='', encoding='latin-1') as file:
                    reader = csv.reader(file)
                    rows = [row for i, row in enumerate(reader) if i < 5]
                    return {
                        "type": "csv",
                        "encoding": "latin-1",
                        "preview": rows
                    }
            except Exception as e:
                return {
                    "type": "csv",
                    "message": f"Error reading CSV file: {str(e)}",
                    "preview": "CSV preview not available"
                }
        except Exception as e:
            return {
                "type": "csv",
                "message": f"Error processing CSV file: {str(e)}",
                "preview": "CSV preview not available"
            }
    
    def generate_html_preview(self, document_data: Dict[str, Any]) -> str:
        """
        Generate HTML preview for a document
        
        Args:
            document_data: Document data from process_document
            
        Returns:
            HTML string for preview
        """
        doc_type = document_data.get("type", "unknown")
        status = document_data.get("status", "error")
        
        if status == "error":
            return f"<div class='error-preview'>Error: {document_data.get('message', 'Unknown error')}</div>"
        
        if doc_type == "pdf":
            return self._generate_pdf_preview(document_data)
        elif doc_type == "docx":
            return self._generate_docx_preview(document_data)
        elif doc_type == "xlsx":
            return self._generate_xlsx_preview(document_data)
        elif doc_type == "csv":
            return self._generate_csv_preview(document_data)
        elif doc_type == "text":
            return self._generate_text_preview(document_data)
        elif doc_type == "visio":
            return self._generate_visio_preview(document_data)
        else:
            return f"<div class='unknown-preview'>No preview available for {doc_type}</div>"
    
    def _generate_pdf_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for PDF"""
        metadata = document_data.get("metadata", {})
        preview = document_data.get("preview", "")
        num_pages = document_data.get("num_pages", 0)
        
        html = f"""
        <div class="pdf-preview">
            <h3>PDF Document</h3>
            <div class="metadata">
                <p><strong>Title:</strong> {metadata.get('title', 'Unknown')}</p>
                <p><strong>Author:</strong> {metadata.get('author', 'Unknown')}</p>
                <p><strong>Pages:</strong> {num_pages}</p>
            </div>
            <div class="content-preview">
                <h4>Content Preview:</h4>
                <pre>{preview}</pre>
            </div>
        </div>
        """
        return html
    
    def _generate_docx_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for DOCX"""
        stats = document_data.get("stats", {})
        preview = document_data.get("preview", "")
        
        html = f"""
        <div class="docx-preview">
            <h3>Word Document</h3>
            <div class="stats">
                <p><strong>Paragraphs:</strong> {stats.get('paragraphs', 0)}</p>
                <p><strong>Tables:</strong> {stats.get('tables', 0)}</p>
                <p><strong>Sections:</strong> {stats.get('sections', 0)}</p>
            </div>
            <div class="content-preview">
                <h4>Content Preview:</h4>
                <pre>{preview}</pre>
            </div>
        </div>
        """
        return html
    
    def _generate_xlsx_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for XLSX"""
        sheets = document_data.get("sheets", [])
        preview = document_data.get("preview", [])
        
        # Generate table HTML
        table_html = "<table class='xlsx-table'>"
        for row in preview:
            table_html += "<tr>"
            for cell in row:
                table_html += f"<td>{cell if cell is not None else ''}</td>"
            table_html += "</tr>"
        table_html += "</table>"
        
        html = f"""
        <div class="xlsx-preview">
            <h3>Excel Spreadsheet</h3>
            <div class="sheets">
                <p><strong>Sheets:</strong> {', '.join(sheets)}</p>
            </div>
            <div class="content-preview">
                <h4>Content Preview:</h4>
                {table_html}
            </div>
        </div>
        """
        return html
    
    def _generate_csv_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for CSV"""
        preview = document_data.get("preview", [])
        
        # Generate table HTML
        table_html = "<table class='csv-table'>"
        for i, row in enumerate(preview):
            table_html += "<tr>"
            for cell in row:
                if i == 0:
                    table_html += f"<th>{cell}</th>"
                else:
                    table_html += f"<td>{cell}</td>"
            table_html += "</tr>"
        table_html += "</table>"
        
        html = f"""
        <div class="csv-preview">
            <h3>CSV Document</h3>
            <div class="content-preview">
                <h4>Content Preview:</h4>
                {table_html}
            </div>
        </div>
        """
        return html
    
    def _generate_text_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for text"""
        stats = document_data.get("stats", {})
        preview = document_data.get("preview", "")
        
        html = f"""
        <div class="text-preview">
            <h3>Text Document</h3>
            <div class="stats">
                <p><strong>Lines:</strong> {stats.get('lines', 0)}</p>
                <p><strong>Words:</strong> {stats.get('words', 0)}</p>
                <p><strong>Characters:</strong> {stats.get('characters', 0)}</p>
            </div>
            <div class="content-preview">
                <h4>Content Preview:</h4>
                <pre>{preview}</pre>
            </div>
        </div>
        """
        return html
    
    def _generate_visio_preview(self, document_data: Dict[str, Any]) -> str:
        """Generate HTML preview for Visio"""
        message = document_data.get("message", "")
        
        html = f"""
        <div class="visio-preview">
            <h3>Visio Document</h3>
            <div class="message">
                <p>{message}</p>
            </div>
        </div>
        """
        return html

# Example usage
if __name__ == "__main__":
    # Configure logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Create document processor
    processor = DocumentProcessor()
    
    # Process a sample document
    import sys
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        result = processor.process_document(file_path)
        print(json.dumps(result, indent=2))
        
        # Generate HTML preview
        html = processor.generate_html_preview(result)
        print("\nHTML Preview:")
        print(html)
