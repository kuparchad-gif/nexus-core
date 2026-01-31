"""
Document Handling Module
Complete toolkit for creating, reading, editing, and managing documents
"""

import os
import shutil
from typing import List, Dict, Any, Optional
from pathlib import Path
import logging

# Document libraries
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import openpyxl
from openpyxl.styles import Font, Fill, PatternFill, Alignment
from openpyxl.chart import BarChart, LineChart, PieChart, Reference
import pandas as pd
from PIL import Image
import json
import csv

logger = logging.getLogger(__name__)


class DocumentHandler:
    """
    Comprehensive document handling for all file types:
    - Text files (.txt, .md, .json, .csv)
    - Word documents (.docx)
    - PDF files (.pdf)
    - Excel spreadsheets (.xlsx)
    - Images (conversion and manipulation)
    """
    
    def __init__(self, workspace: str = "/tmp/documents"):
        self.workspace = Path(workspace)
        self.workspace.mkdir(parents=True, exist_ok=True)
        
    # ==================== TEXT FILES ====================
    
    def create_text_file(self, filename: str, content: str, encoding: str = 'utf-8') -> str:
        """Create or overwrite a plain text file"""
        filepath = self.workspace / filename
        try:
            with open(filepath, 'w', encoding=encoding) as f:
                f.write(content)
            logger.info(f"Text file created: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"Error creating text file: {e}")
            raise
    
    def read_text_file(self, filename: str, encoding: str = 'utf-8') -> str:
        """Read content from a text file"""
        filepath = self.workspace / filename
        try:
            with open(filepath, 'r', encoding=encoding) as f:
                content = f.read()
            return content
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            raise
    
    def append_to_text_file(self, filename: str, content: str, encoding: str = 'utf-8') -> str:
        """Append content to existing text file"""
        filepath = self.workspace / filename
        try:
            with open(filepath, 'a', encoding=encoding) as f:
                f.write(content)
            logger.info(f"Content appended to: {filepath}")
            return str(filepath)
        except Exception as e:
            logger.error(f"Error appending to text file: {e}")
            raise
    
    # ==================== WORD DOCUMENTS ====================
    
    def create_word_document(self, filename: str, title: str, content: List[Dict[str, Any]]) -> str:
        """
        Create a Word document with rich formatting
        content format: [
            {"type": "heading", "text": "Title", "level": 1},
            {"type": "paragraph", "text": "Content", "bold": False, "italic": False},
            {"type": "bullet", "items": ["item1", "item2"]},
            {"type": "table", "data": [[row1], [row2]], "headers": ["col1", "col2"]},
            {"type": "image", "path": "/path/to/image.jpg", "width": 6}
        ]
        """
        filepath = self.workspace / filename
        doc = docx.Document()
        
        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process content
        for item in content:
            item_type = item.get("type", "paragraph")
            
            if item_type == "heading":
                doc.add_heading(item["text"], level=item.get("level", 1))
            
            elif item_type == "paragraph":
                para = doc.add_paragraph(item["text"])
                if item.get("bold"):
                    para.runs[0].bold = True
                if item.get("italic"):
                    para.runs[0].italic = True
                if item.get("font_size"):
                    para.runs[0].font.size = Pt(item["font_size"])
            
            elif item_type == "bullet":
                for bullet_item in item["items"]:
                    doc.add_paragraph(bullet_item, style='List Bullet')
            
            elif item_type == "numbered":
                for num_item in item["items"]:
                    doc.add_paragraph(num_item, style='List Number')
            
            elif item_type == "table":
                data = item["data"]
                headers = item.get("headers", [])
                
                table = doc.add_table(rows=len(data) + (1 if headers else 0), cols=len(data[0]))
                table.style = 'Light Grid Accent 1'
                
                if headers:
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True
                    row_offset = 1
                else:
                    row_offset = 0
                
                for i, row in enumerate(data):
                    for j, cell_value in enumerate(row):
                        table.rows[i + row_offset].cells[j].text = str(cell_value)
            
            elif item_type == "image":
                if os.path.exists(item["path"]):
                    doc.add_picture(item["path"], width=Inches(item.get("width", 6)))
        
        doc.save(str(filepath))
        logger.info(f"Word document created: {filepath}")
        return str(filepath)
    
    def read_word_document(self, filename: str) -> Dict[str, Any]:
        """Extract text and structure from Word document"""
        filepath = self.workspace / filename
        doc = docx.Document(str(filepath))
        
        content = {
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }
        
        for para in doc.paragraphs:
            content["paragraphs"].append({
                "text": para.text,
                "style": para.style.name
            })
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            content["tables"].append(table_data)
        
        content["full_text"] = "\n".join([p.text for p in doc.paragraphs])
        
        return content
    
    # ==================== PDF FILES ====================
    
    def read_pdf(self, filename: str) -> Dict[str, Any]:
        """Extract text from PDF with page information"""
        filepath = self.workspace / filename
        reader = PdfReader(str(filepath))
        
        content = {
            "num_pages": len(reader.pages),
            "pages": [],
            "full_text": "",
            "metadata": reader.metadata
        }
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            content["pages"].append({
                "page_num": i + 1,
                "text": page_text
            })
            content["full_text"] += page_text + "\n"
        
        return content
    
    def merge_pdfs(self, pdf_files: List[str], output_filename: str) -> str:
        """Merge multiple PDF files into one"""
        merger = PdfMerger()
        
        for pdf_file in pdf_files:
            filepath = self.workspace / pdf_file
            merger.append(str(filepath))
        
        output_path = self.workspace / output_filename
        merger.write(str(output_path))
        merger.close()
        
        logger.info(f"PDFs merged: {output_path}")
        return str(output_path)
    
    def split_pdf(self, filename: str, page_ranges: List[tuple]) -> List[str]:
        """
        Split PDF into multiple files based on page ranges
        page_ranges: [(1, 5), (6, 10), ...] - inclusive ranges
        """
        filepath = self.workspace / filename
        reader = PdfReader(str(filepath))
        output_files = []
        
        for i, (start, end) in enumerate(page_ranges):
            writer = PdfWriter()
            
            for page_num in range(start - 1, end):
                if page_num < len(reader.pages):
                    writer.add_page(reader.pages[page_num])
            
            output_filename = f"{Path(filename).stem}_part{i+1}.pdf"
            output_path = self.workspace / output_filename
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            output_files.append(str(output_path))
        
        logger.info(f"PDF split into {len(output_files)} files")
        return output_files
    
    def extract_pdf_pages(self, filename: str, page_numbers: List[int], output_filename: str) -> str:
        """Extract specific pages from PDF"""
        filepath = self.workspace / filename
        reader = PdfReader(str(filepath))
        writer = PdfWriter()
        
        for page_num in page_numbers:
            if 0 < page_num <= len(reader.pages):
                writer.add_page(reader.pages[page_num - 1])
        
        output_path = self.workspace / output_filename
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        logger.info(f"Extracted pages to: {output_path}")
        return str(output_path)
    
    # ==================== EXCEL SPREADSHEETS ====================
    
    def create_excel(self, filename: str, sheets: Dict[str, List[List[Any]]],
                     headers: Optional[Dict[str, List[str]]] = None,
                     charts: Optional[Dict[str, Dict]] = None) -> str:
        """
        Create Excel file with multiple sheets, headers, and charts
        sheets: {"Sheet1": [[row1], [row2], ...], "Sheet2": ...}
        headers: {"Sheet1": ["Col1", "Col2"], ...}
        charts: {"Sheet1": {"type": "bar", "range": "A1:B10", "title": "Chart"}}
        """
        filepath = self.workspace / filename
        wb = openpyxl.Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        for sheet_name, data in sheets.items():
            ws = wb.create_sheet(sheet_name)
            
            # Add headers if provided
            if headers and sheet_name in headers:
                ws.append(headers[sheet_name])
                # Style headers
                for cell in ws[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                    cell.alignment = Alignment(horizontal="center")
            
            # Add data
            for row in data:
                ws.append(row)
            
            # Add charts if specified
            if charts and sheet_name in charts:
                chart_config = charts[sheet_name]
                chart_type = chart_config.get("type", "bar")
                
                if chart_type == "bar":
                    chart = BarChart()
                elif chart_type == "line":
                    chart = LineChart()
                elif chart_type == "pie":
                    chart = PieChart()
                else:
                    chart = BarChart()
                
                chart.title = chart_config.get("title", "Chart")
                chart.style = 10
                
                # Add data to chart
                data_range = chart_config.get("range", "A1:B10")
                data_ref = Reference(ws, range_string=data_range)
                chart.add_data(data_ref, titles_from_data=True)
                
                ws.add_chart(chart, chart_config.get("position", "E5"))
        
        wb.save(str(filepath))
        logger.info(f"Excel file created: {filepath}")
        return str(filepath)
    
    def read_excel(self, filename: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """Read Excel file and return data from all or specific sheet"""
        filepath = self.workspace / filename
        wb = openpyxl.load_workbook(str(filepath))
        
        result = {"sheets": {}}
        
        sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames
        
        for sheet in sheets_to_read:
            ws = wb[sheet]
            data = []
            for row in ws.iter_rows(values_only=True):
                data.append(list(row))
            result["sheets"][sheet] = data
        
        return result
    
    def edit_excel_cell(self, filename: str, sheet_name: str, cell: str, value: Any) -> str:
        """Edit a specific cell in Excel file"""
        filepath = self.workspace / filename
        wb = openpyxl.load_workbook(str(filepath))
        ws = wb[sheet_name]
        ws[cell] = value
        wb.save(str(filepath))
        logger.info(f"Cell {cell} updated in {filepath}")
        return str(filepath)
    
    def excel_to_dataframe(self, filename: str, sheet_name: str = None) -> pd.DataFrame:
        """Convert Excel sheet to pandas DataFrame"""
        filepath = self.workspace / filename
        if sheet_name:
            df = pd.read_excel(filepath, sheet_name=sheet_name)
        else:
            df = pd.read_excel(filepath)
        return df
    
    def dataframe_to_excel(self, df: pd.DataFrame, filename: str, sheet_name: str = "Sheet1") -> str:
        """Save pandas DataFrame to Excel"""
        filepath = self.workspace / filename
        df.to_excel(filepath, sheet_name=sheet_name, index=False)
        logger.info(f"DataFrame saved to Excel: {filepath}")
        return str(filepath)
    
    # ==================== CSV FILES ====================
    
    def create_csv(self, filename: str, data: List[List[Any]], headers: Optional[List[str]] = None) -> str:
        """Create CSV file"""
        filepath = self.workspace / filename
        
        with open(filepath, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            if headers:
                writer.writerow(headers)
            writer.writerows(data)
        
        logger.info(f"CSV file created: {filepath}")
        return str(filepath)
    
    def read_csv(self, filename: str) -> Dict[str, Any]:
        """Read CSV file"""
        filepath = self.workspace / filename
        
        with open(filepath, 'r', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            data = list(reader)
        
        return {
            "headers": data[0] if data else [],
            "data": data[1:] if len(data) > 1 else [],
            "full_data": data
        }
    
    # ==================== JSON FILES ====================
    
    def create_json(self, filename: str, data: Dict[str, Any], indent: int = 2) -> str:
        """Create JSON file"""
        filepath = self.workspace / filename
        
        with open(filepath, 'w', encoding='utf-8') as jsonfile:
            json.dump(data, jsonfile, indent=indent, ensure_ascii=False)
        
        logger.info(f"JSON file created: {filepath}")
        return str(filepath)
    
    def read_json(self, filename: str) -> Dict[str, Any]:
        """Read JSON file"""
        filepath = self.workspace / filename
        
        with open(filepath, 'r', encoding='utf-8') as jsonfile:
            data = json.load(jsonfile)
        
        return data
    
    # ==================== FILE OPERATIONS ====================
    
    def copy_file(self, source: str, destination: str) -> str:
        """Copy file to new location"""
        src_path = self.workspace / source
        dst_path = self.workspace / destination
        shutil.copy2(src_path, dst_path)
        logger.info(f"File copied: {src_path} -> {dst_path}")
        return str(dst_path)
    
    def move_file(self, source: str, destination: str) -> str:
        """Move file to new location"""
        src_path = self.workspace / source
        dst_path = self.workspace / destination
        shutil.move(str(src_path), str(dst_path))
        logger.info(f"File moved: {src_path} -> {dst_path}")
        return str(dst_path)
    
    def delete_file(self, filename: str) -> bool:
        """Delete file"""
        filepath = self.workspace / filename
        try:
            os.remove(filepath)
            logger.info(f"File deleted: {filepath}")
            return True
        except Exception as e:
            logger.error(f"Error deleting file: {e}")
            return False
    
    def list_files(self, pattern: str = "*") -> List[str]:
        """List files in workspace matching pattern"""
        files = list(self.workspace.glob(pattern))
        return [str(f.name) for f in files if f.is_file()]
    
    def file_info(self, filename: str) -> Dict[str, Any]:
        """Get file information"""
        filepath = self.workspace / filename
        stat = filepath.stat()
        
        return {
            "name": filename,
            "path": str(filepath),
            "size_bytes": stat.st_size,
            "size_mb": stat.st_size / (1024 * 1024),
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "extension": filepath.suffix
        }
    
    # ==================== CONVERSION ====================
    
    def convert_format(self, filename: str, target_format: str) -> str:
        """
        Convert document between formats
        Supports: txt, docx, pdf, csv, json, xlsx
        """
        source_path = self.workspace / filename
        source_ext = source_path.suffix.lower()
        target_ext = f".{target_format.lower().strip('.')}"
        output_filename = source_path.stem + target_ext
        output_path = self.workspace / output_filename
        
        # Text to other formats
        if source_ext == '.txt':
            content = self.read_text_file(filename)
            
            if target_ext == '.docx':
                doc = docx.Document()
                for line in content.split('\n'):
                    doc.add_paragraph(line)
                doc.save(str(output_path))
            
            elif target_ext == '.json':
                data = {"content": content}
                self.create_json(output_filename, data)
        
        # Excel to CSV
        elif source_ext == '.xlsx' and target_ext == '.csv':
            df = self.excel_to_dataframe(filename)
            df.to_csv(output_path, index=False)
        
        # CSV to Excel
        elif source_ext == '.csv' and target_ext == '.xlsx':
            df = pd.read_csv(source_path)
            self.dataframe_to_excel(df, output_filename)
        
        logger.info(f"File converted: {filename} -> {output_filename}")
        return str(output_path)
